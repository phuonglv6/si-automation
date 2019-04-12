from __future__ import division
from .pdfMatching import matchingForm, handleMatchedFile, handleUnmatchedFile
from docx import Document
import sys
import re
import subprocess
import os
from django.conf import settings

# import time
# import sys

# file_path = sys.argv[1]
# outPath = sys.argv[2]
keys = settings.SIBL_CLASSES.values()


def extract_word(file_path):
    data_json = {}
    index_key = []
    document = Document(file_path)
    tables = document.tables
    key_value = []
    symbols = ['', ' ']
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                data = []
                for paragraph in cell.paragraphs:
                    if paragraph.text not in data:
                        data.append(paragraph.text)
                for symbol in symbols:
                    while symbol in data:
                        data.remove(symbol)

                if data not in key_value:
                    key_value.append(data)

    for key in key_value:
        if len(key) > 0:
            array = []
            for k in keys:
                if k in key[0].upper():
                    array.append(1)
                else:
                    array.append(0)
            for index, val in enumerate(array):
                if val == 1:
                    index_key.append(index)
                    value = ''
                    for i in range(1, len(key)):
                        if value == '':
                            value = value + key[i]
                        else:
                            value = value + '\n' + key[i]
                    data_json[keys[index]] = value
        if len(key) == 1 and 'kgs'.upper() in key[0].upper():
            data_json['total_gross_weight'] = key[0]
    for i in range(len(keys)):
        if i not in index_key:
            data_json[keys[i]] = ''

    return data_json


class LibreOfficeError(Exception):
    def __init__(self, output):
        self.output = output


def libreoffice_exec():
    if sys.platform == 'darwin':
        return '/Applications/LibreOffice.app/Contents/MacOS/soffice'
    elif sys.platform == 'win32':
        return 'C:/Program Files/LibreOffice/program/soffice'
    elif 'linux' in sys.platform:
        return '/usr/bin/soffice'
    else:
        print('--------ERROR--------')
        print('Unknown Platform')
        return 'libreoffice'


def convert_to(folder, source, timeout=None):
    args = [libreoffice_exec(), '--headless', '--convert-to',
            'pdf', '--outdir', folder, source]

    process = subprocess.run(args, stdout=subprocess.PIPE,
                             stderr=subprocess.PIPE, timeout=timeout)
    filename = re.search('-> (.*?) using filter', process.stdout.decode())

    if filename is None:
        raise LibreOfficeError(process.stdout.decode())
    else:
        return filename.group(1)


def extract_text(file_path, outPath):
    pdf_path = file_type_to_pdf(file_path)
    chosenFolder, img, num_pages = matchingForm(pdf_path, outPath)
    if chosenFolder:  # Match template in DB
        jsonName = pdf_path.split('/')[-1].split('.')[0] + '.json'
        data_result = handleMatchedFile(
            chosenFolder, outPath, img, jsonName, pdf_path)
    else:  # new type case
        data_result = handleUnmatchedFile(outPath, img)

    return data_result, num_pages


def file_type_to_pdf(file_path):
    if file_path.endswith('.docx') or file_path.endswith('.doc') or \
            file_path.endswith('.xls') or file_path.endswith('.xlsx'):
        pdf_path = convert_to(settings.MEDIA_ROOT, file_path)
        os.remove(file_path)
        return pdf_path
    else:
        return file_path


# def handle_pdf(file_path, outPath):
#     chosenFolder, img, num_pages = matchingForm(file_path, outPath)
#     if chosenFolder:  # Match template in DB
#         num_pages = None
#         jsonName = file_path.split('/')[-1].split('.')[0] + '.json'
#         data_result = handleMatchedFile(
#                 chosenFolder, outPath, img, jsonName, file_path)
#     else:  # new type case
#         data_result = handleUnmatchedFile(outPath, img)

#     return data_result, num_pages


# def extract_text(file_path, outPath):
#     if file_path.endswith('.docx') or file_path.endswith('.doc') or \
#           file_path.endswith('.xls') or file_path.endswith('.xlsx'):
#         filename = convert_to(outPath, file_path)
#         data_result, num_pages = handle_pdf(filename, outPath)
#         os.remove(filename)
#     else:
#         data_result, num_pages = handle_pdf(file_path, outPath)
#     return data_result, num_pages
