from __future__ import division
import sys
import re
import subprocess
import os
from docx import Document
from flask import current_app
from ..import db
from .pdf_matching import matchingForm, handleMatchedFile, handleUnmatchedFile
from ..util.consts import SIBL_CLASSES
from ..api.models import Doc, ContainerDetail
from .get_container_detail import extract_container_detail
from .get_doc_detail import extract_vessel_voyage, extract_bkg_no, \
    clean_marks, hs_code_process, recheck_total_packages, \
    recheck_conts_package, clean_unit_slash, totalize_measurement, \
    split_por_del  # clean_content_from

keys = SIBL_CLASSES.values()
classes = SIBL_CLASSES

def extract_word(file_path):
    data_json = {}
    index_key = []
    document = Document(file_path)
    tables = document.tables
    key_value = []
    symbols = ['', ' ']
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                data = []
                for paragraph in cell.paragraphs:
                    if paragraph.text not in data:
                        data.append(paragraph.text)
                for symbol in symbols:
                    while symbol in data:
                        data.remove(symbol)

                if data not in key_value:
                    key_value.append(data)

    for key in key_value:
        if len(key) > 0:
            array = []
            for k in keys:
                if k in key[0].upper():
                    array.append(1)
                else:
                    array.append(0)
            for index, val in enumerate(array):
                if val == 1:
                    index_key.append(index)
                    value = ''
                    for i in range(1, len(key)):
                        if value == '':
                            value = value + key[i]
                        else:
                            value = value + '\n' + key[i]
                    data_json[keys[index]] = value
        if len(key) == 1 and 'kgs'.upper() in key[0].upper():
            data_json['total_gross_weight'] = key[0]
    for i in range(len(keys)):
        if i not in index_key:
            data_json[keys[i]] = ''

    return data_json


class LibreOfficeError(Exception):
    def __init__(self, output):
        self.output = output


def libreoffice_exec():
    if sys.platform == 'darwin':
        return '/Applications/LibreOffice.app/Contents/MacOS/soffice'
    elif sys.platform == 'win32':
        return 'C:/Program Files/LibreOffice/program/soffice'
    elif 'linux' in sys.platform:
        return '/usr/bin/soffice'
    else:
        print('--------ERROR--------')
        print('Unknown Platform')
        return 'libreoffice'


def convert_to(folder, source, timeout=None):
    args = [libreoffice_exec(), '--headless', '--convert-to',
            'pdf', '--outdir', folder, source]

    process = subprocess.run(args, stdout=subprocess.PIPE,
                             stderr=subprocess.PIPE, timeout=timeout)
    filename = re.search('-> (.*?) using filter', process.stdout.decode())

    if filename is None:
        raise LibreOfficeError(process.stdout.decode())
    else:
        return filename.group(1)


def extract_text(file_path, outPath):
    pdf_path = file_type_to_pdf(file_path)
    print("file_type_to_pdf: ",pdf_path)
    chosenNo, img, num_pages, sameTemplateNo = matchingForm(pdf_path, outPath)
    if chosenNo:  # Match template in DB
        jsonName = pdf_path.split('/')[-1].split('.')[0] + '.json'
        data_result = handleMatchedFile(
            chosenNo, outPath, img, jsonName, pdf_path)
    else:  # new type case
        data_result = handleUnmatchedFile(outPath, img, sameTemplateNo)
    print("extract_text: data_result ",data_result)
    return data_result, num_pages,chosenNo


def file_type_to_pdf(file_path):
    if file_path.endswith('.docx') or file_path.endswith('.doc') or \
            file_path.endswith('.xls') or file_path.endswith('.xlsx'):
        pdf_path = convert_to(current_app.config['UPLOAD_FOLDER'], file_path)
        os.remove(file_path)
        return pdf_path
    else:
        return file_path

def extract_file(filename):
    """ Get:    uploaded filename and forward to extracting
        Return: status and data dict
    """
    file_path = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)
    out_path = current_app.config['OUT_EXTRACTION_FOLDER']
    print("extract_file: file_path: ",file_path)
    print("extract_file: out_path: ",out_path)
    si_result, num_pages, chosenNo = extract_text(file_path, out_path)
    # print(si_result)

    data = {}
    status = 200
    if isinstance(si_result, dict):
        si_model = {
            "name": filename,
            "template": chosenNo
            }
        for v in classes.values():
            if v not in list(classes.values())[20:]:
                si_model[v] = si_result.get(v, '')

        # TODO: Re-factor SI_classes in settings
        si_model['total_type'] = si_result.get('total_type', '')
        si_model['also_notify'] = si_result.get('also_notify', '')
        si_model['final_destination'] = si_result.get('final_destination', '')

        # Split Vessel Voyage
        vessel = si_result.get('vessel', '')
        voyage = si_result.get('voyage', '')
        doNothing, vessel, voyage = extract_vessel_voyage(vessel, voyage)
        if doNothing is False:
            print('SPLITING VESSEL - VOYAGE')
            print(vessel, ' --- ', voyage)
            si_model['vessel'] = vessel
            si_model['voyage'] = voyage

        # HS CODE
        hs_codes, bl_type = hs_code_process(si_result)
        print('EXTRACTING HS CODE - BL Type')
        print('HS CODE:', hs_codes)
        print('BL Type:', bl_type)
        si_model['hs_code'] = ', '.join(hs_codes)
        si_model['bl_type'] = bl_type

        # MARKS
        doNothing, conts_detail, total_mark, bkg_no = clean_marks(si_result)
        if doNothing is False:
            print('SPLITING BOOKING NUMBER - MARKS - CONT. DETAIL ')
            print('BOOKING NUMBER:', bkg_no)
            print('MARKS:', total_mark)
            print('CONT. DETAIL:', conts_detail)
            si_model['total_mark'] = total_mark
            si_model['bkg_no'] = bkg_no
            si_result['containers_detail'] = conts_detail

        print('MARKS DONE')

        # BOOKING NUMBER have noise inside
        si_model['bkg_no'] = extract_bkg_no(
            si_model['bkg_no'],
            si_result.get('so_number', '')
            )
        print('BKG NO DONE')
        print("+++++++++++++++++++++++si_model: ",si_model)
        # CLEAN UNIT Measurement noise char
        for unit in ['total_packages', 'total_weight', 'total_measurement']:
            si_model[unit] = clean_unit_slash(si_model[unit])

        # VN 105604 have inline POR and DEL
        if si_model['place_of_receipt'] == si_model['place_of_delivery'] \
                and si_model['port_of_discharge'] \
                and si_model['place_of_delivery']:
            por, podel = split_por_del(si_model['place_of_receipt'],
                                       si_model['port_of_discharge'])
            if por:
                si_model['place_of_receipt'] = por
            if podel:
                si_model['place_of_delivery'] = podel

        # Save Doc
        # doc = Doc.objects.create(**si_model)
        # data['doc'] = model_to_dict(doc)
        doc = Doc(**si_model)
        db.session.add(doc)
        db.session.commit()
        data['doc'] = doc.to_dict(show_all=True)
        # TODO: How about Cont. Number, Cont. Seal Annotation for VN102347
        # Continue with Doc's Container
        if 'containers_detail' in si_result:
            cont_dicts, is_verified = extract_container_detail(si_result)
            cont_dicts = recheck_conts_package(cont_dicts, si_result)
            pkgs_sum, checked = recheck_total_packages(cont_dicts, si_result)
            if checked:
                doc_temp = Doc.query.filter_by(id=doc.id).first()
                doc_temp.total_packages = pkgs_sum
            db.session.commit()
            si_model = totalize_measurement(cont_dicts, doc_temp.to_dict(show_all=True))

            data['conts'] = []
            for detail in cont_dicts:
                for unit in ['packages', 'container_weight', 'container_measurement']:
                    detail[unit] = clean_unit_slash(detail.get(unit, ''))

                detail['doc'] = doc
                # cont = ContainerDetail.objects.create(**detail)
                cont = ContainerDetail(**detail)
                # data['conts'].append(model_to_dict(cont))
                data['conts'].append(cont.to_dict())

            data['is_verified'] = is_verified
            
            status = 200
    else:
        try:
            temp_number = int(si_result)
        except TypeError:
            # extract_text not return Dict or Template Number
            data['error'] = 'extract result is empty'
            status = 500
            print('Does not receive Template Number')
        else:
            status = 200
            data['temp_num'] = temp_number

    return status, data, num_pages